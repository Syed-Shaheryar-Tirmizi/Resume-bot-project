"""Build a formatted Word resume from chat-style markdown text."""

import re
from io import BytesIO

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING


def _add_paragraph_runs(paragraph, line: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", line)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run.text = part[2:-2]
            run.bold = True
        else:
            run.text = part


def build_resume_docx(text: str) -> BytesIO:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    style.paragraph_format.space_after = Pt(6)

    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("###"):
            p = doc.add_heading(stripped.lstrip("#").strip(), level=3)
            p.paragraph_format.space_after = Pt(4)
        elif stripped.startswith("##"):
            p = doc.add_heading(stripped.lstrip("#").strip(), level=2)
            p.paragraph_format.space_after = Pt(6)
        elif stripped.startswith("#"):
            p = doc.add_heading(stripped.lstrip("#").strip(), level=1)
            p.paragraph_format.space_after = Pt(8)
        elif re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            body = re.sub(r"^[-*]\s+", "", stripped)
            _add_paragraph_runs(p, body)
        else:
            p = doc.add_paragraph()
            _add_paragraph_runs(p, stripped)

        i += 1

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
